from pathlib import Path

from docx import Document
import os
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_COLOR


from config.settings import UPLOAD_FOLDER


def extract_t_scores(doc_path, table_indices):
    """
    Extract labels and T-scores from specific tables within a document.
    :param doc_path: Path to the Word document.
    :param table_indices: List of table indices to extract data from.
    :return: A dictionary with labels as keys and T-scores as values.
    """
    doc = Document(doc_path)
    t_scores = {}
    print(table_indices)
    for index in table_indices:
        try:
            table = doc.tables[index]
        except IndexError:
            print(f"Table index {index} is out of range for {doc_path}.")
            continue
        print(index)
        for row in table.rows[1:]:  # Skip header row
            cells = row.cells
            label = cells[0].text.strip()
            print(label)


            t_score = cells[2].text.strip()  # Corrected to get T-Score from the 3rd column
            t_scores[label] = t_score

    return t_scores


def update_template(template_path, source_data, prs_trs_srp, reportFile):
    """
    Update the template document with T-scores from the source data.
    :param template_path: Path to the template Word document.
    :param source_data: Dictionary with labels as keys and T-scores as values.
    """
    if prs_trs_srp == 'prs1':
        table_number = 5
        column_number = 3
    elif prs_trs_srp == 'prs2':
        table_number = 5
        column_number = 4
    elif prs_trs_srp == 'trs':
        table_number = 5
        column_number = 5
    elif prs_trs_srp == 'srp':
        table_number = 6
        column_number = 2
    else:
        print('Invalid prs_trs_srp')
        return

    doc = Document(template_path)
    table = doc.tables[table_number]  # Assuming the 4th table is the target

    for row in table.rows[1:]:
        cells = row.cells
        label = cells[1].text.strip()
        if label == '':
            continue
        if label in source_data:
            cell = cells[column_number]
            # Clear previous text
            cell.text = ""
            para = cells[column_number].paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center align text horizontally
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER  # Center align text vertically

            score = source_data[label]
            if label in ['Adaptive Skills', 'Adaptability', 'Social Skills', 'Leadership', 'Activities of Daily Living', 'Functional Communication']:
                if score.isdigit() and int(score) <= 40:
                    # Formatting for score <= 40
                    run = para.add_run(score + '*')
                    run.font.highlight_color = WD_COLOR.YELLOW #RGBColor(255, 0, 0)  # Set font color to red
                elif score.isdigit() and int(score) <= 30:
                    # Formatting for score <= 30
                    run = para.add_run(score + '**')
                    run.font.highlight_color = WD_COLOR.RED #RGBColor(255, 0, 0)  # Set font color to red
                else:
                    run = para.add_run(score)
            else:
                if score.isdigit() and int(score) >= 70:
                    # Formatting for score >= 70
                    run = para.add_run(score + '**')
                    run.font.highlight_color = WD_COLOR.RED #RGBColor(255, 0, 0)  # Set font color to red
                elif score.isdigit() and int(score) >= 60:
                    # Formatting for score >= 60
                    run = para.add_run(score + '*')
                    run.font.highlight_color = WD_COLOR.YELLOW #RGBColor(255, 0, 0)  # Set font color to red
                else:
                    run = para.add_run(score)
        else:
            # Handle case where label is not in source_data
            cell = cells[column_number]
            cell.text = "--"
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            # Save the updated document
    #doc.save(os.path.join(UPLOAD_FOLDER, reportFile))
    #doc.save(os.path.join('', reportFile))
    if os.name == 'posix':  # Unix/Linux/MacOS/Ubuntu
        save_path = doc.save(os.path.join(UPLOAD_FOLDER, reportFile))
    else:  # Windows, os.name will be 'nt'
        save_path = doc.save(reportFile)

# Set the paths for the source and template documents
#prs_doc_path = 'bascprs1.docx'
#srp_doc_path = 'bascsrp.docx'
#template_doc_path = 'template.docx'

#if prs_doc_path:
    # Extract T-scores from the 4th and 5th tables
#    t_scores = extract_t_scores(prs_doc_path, [3, 5, 7, 8])  # Tables are 0-indexed
#    print(t_scores)
    # Update the template with the extracted T-scores
#    update_template(template_doc_path, t_scores, 'prs')

#if srp_doc_path:
#    t_scores = extract_t_scores(srp_doc_path, [3, 5, 7])  # Tables are 0-indexed
#    update_template('updated_template.docx', t_scores, 'srp')
