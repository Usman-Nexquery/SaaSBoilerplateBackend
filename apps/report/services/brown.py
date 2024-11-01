import os
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import fitz  # PyMuPDF

from config.settings import UPLOAD_FOLDER

from .hi import get_indices

def set_cell_vertical_alignment(cell, align="center"):
    """Set the vertical alignment of text within a cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    valign = OxmlElement('w:vAlign')
    valign.set(qn('w:val'), align)  # 'top', 'center', or 'bottom'
    tcPr.append(valign)


def is_convertible_to_int(val):
    try:
        int(val)  # Try converting to integer
        return True
    except ValueError:
        return False
def extract_score_summary_table(text):
    known_labels = ["Activation", "Focus", "Effort", "Emotion", "Memory", "Action", "Total Composite"]
    persistent_list = []  # List to store final blocks of data
    temp_list = []  # Temporary list to store data between known labels
    label_found = False  # Flag to indicate if we are currently capturing rows
    row_count_since_last_label = 0  # Counter to track rows since the last label

    rows = text.split('\n')  # Assuming text is a single string that needs to be split into rows

# 60-64: At Risk

    for row in rows:

        row_trimmed = row.strip()  # Trim whitespace to ensure accurate matching
        if any(row_trimmed.startswith(label) for label in known_labels) and any(row_trimmed == label for label in known_labels):  # Check if the row starts with a known label
            if temp_list:  # Append current list to persistent if it's not empty
                if len(temp_list) > 3: persistent_list.append(temp_list)

            temp_list = [row]  # Start a new list with the current label
            label_found = True
            row_count_since_last_label = 0
        elif label_found:
            temp_list.append(row)  # Append current row to temporary list
            row_count_since_last_label += 1
            if row_count_since_last_label > 4:  # Stop appending if 8 rows pass without a new label
                if temp_list:
                    if any(is_convertible_to_int(val) for val in temp_list): persistent_list.append(temp_list)
                temp_list = []
                label_found = False
                row_count_since_last_label = 0

    if temp_list:  # Check at the end to append any remaining data
        if any(is_convertible_to_int(val) for val in temp_list): persistent_list.append(temp_list)

    return persistent_list


def extract_data_from_pdf(pdf_path):
    doc = fitz.open(pdf_path)
    all_tables = []  # Store tables from all pages
    for page in doc:
        text = page.get_text("text")
        # Extract the specific table data
        table_data = extract_score_summary_table(text)
        if table_data:
            all_tables.append(table_data)

            doc.close()  # Close the document
            return all_tables  # Stop further processing and return the tables
        else:
            print("No data found on this page.")

    doc.close()  # Ensure the document is closed if no tables are found
    return all_tables

def insert_table_into_word(doc_path, data, filename):
    # Load the Word document
    doc = Document(doc_path)
    table_indices = get_indices(doc, "Brown")
    # print(table_indices)
    brown_index = [3, 4]


    # Process each table found by the getIndices function
    for table_name, info in table_indices.items():
        table = doc.tables[info["table_index"]]  # Access the specific table by index
        column_indices = info["columns"]  # Get the list of column indices needed
        # column_indices = [x + 1 for x in column_indices]
        # Iterate over each row in the table after the header row
        for row in table.rows[1:]:  # Skipping the header row
            key = row.cells[0].text.strip()
            #if key in data:
            if any(key in sublist for sublist in data):
                # Insert data into specified columns
                for k in range(len(data)):
                    for i, column_index in enumerate(column_indices):
                        cell = row.cells[column_index]
                        if data[k][0] == key:
                            cell.text = str(data[k][brown_index[i]])  # Assuming data[key] is a list with the needed data

                        paragraph = cell.paragraphs[0]
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        # Assume set_cell_vertical_alignment is defined elsewhere
                        set_cell_vertical_alignment(cell, "center")

    if os.name == 'posix':  # Unix/Linux/MacOS/Ubuntu
        doc.save(os.path.join(UPLOAD_FOLDER, filename))
    else:  # Windows, os.name will be 'nt'
        doc.save(filename)
def generate_document_text(tables):
    significant_count = 0
    total_clusters = 6  # Assuming there are 6 clusters as mentioned
    for table in tables:
        # Assuming the fourth element in each sublist of the table is the T-score
        t_score = int(table[3])  # Convert T-score to integer
        if t_score >= 75: significant_count += 1

    # Generate the text based on the count of significant scores
    if significant_count == total_clusters:
        report_text = f"[First Name] reported clinically significant scores about [Him/Her]self on all 6 clusters."
    else:
        report_text = f"[First Name] reported clinically significant scores about [Him/Her]self on {significant_count} clusters."

    return report_text

def Brown(file):
    dirname = os.path.dirname(__file__)
    pdf_path = os.path.join(dirname, 'files/Brown.pdf')
    doc_path = os.path.join(dirname, 'testing.docx')
    tables = extract_data_from_pdf(pdf_path)
    filename = 'testing.docx'
    if tables:
        for table in tables:
            print("Extracted Table:", table)
    else:
        print("No tables extracted from the document.")


    insert_table_into_word(file, tables[0], filename)

    brownText = generate_document_text(tables[0][:-1])

    return brownText
    #insert_table_into_word(doc_path, tables, filename)


#if __name__ == "__main__":
#    Brown('/home/ubuntu/flask/files/ADHD.docx')
