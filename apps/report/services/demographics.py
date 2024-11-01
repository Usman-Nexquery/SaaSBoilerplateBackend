from pathlib import Path

from docx import Document
import os

from config.settings import UPLOAD_FOLDER


def extract_examinee_info(doc_path):
    doc = Document(doc_path)
    examinee_info = {}

    for table in doc.tables:
        first_cell_text = table.cell(0, 0).text.strip()

        if "Examinee Name" in first_cell_text or "Date of Testing" in first_cell_text:  # Adjusted to include both tables
            for row in table.rows:
                cell_texts = [cell.text.strip() for cell in row.cells]

                # Extract pairs based on the expected layout
                for i in range(0, len(row.cells), 2):  # Handles pairs (step by 2)
                    if i + 1 < len(row.cells):  # Check if there's a corresponding value cell
                        key = row.cells[i].text.strip()
                        value = row.cells[i+1].text.strip()
                        if key:  # Only add to dict if key exists

                            examinee_info[key] = value
                            if "Date of Testing" in key:
                                key = row.cells[3].text.strip()
                                value = row.cells[4].text.strip()
                                examinee_info[key] = value
                                break

    replacements = {}
    replacements['[Full Name]'] = examinee_info['Examinee Name']
    replacements['[First Name]'] = examinee_info['Examinee Name'].split(" ")[0]
    replacements['[Last Name]'] = examinee_info['Examinee Name'].split(" ")[1]
    replacements['[Exact Age]'] = examinee_info['Age at Testing']
    replacements['[Eval Date]'] = examinee_info['Date of Testing']
    replacements['[DOB]'] = examinee_info['Date of Birth']
    replacements['[Age]'] = examinee_info['Age at Testing'].split(" ")[0]

    return replacements
def replace_placeholders(doc_path, firstName, pronouns, brownText):
    # Open the document
    doc = Document(doc_path)

    # Define the replacements based on pronouns
    if pronouns.lower() == "male":
        pronoun_replacements = {
            "[He/She]": "He",
            "[His/Her]": "his",
            "[Him/Her]": "him",
            "[he/she]": "he",
            "[Son/Daughter]": "daughter"
        }
    elif pronouns.lower() == "female":
        pronoun_replacements = {
            "[He/She]": "She",
            "[His/Her]": "her",
            "[Him/Her]": "her",
            "[he/she]": "she",
            "[Son/Daughter]": "daughter"
        }
    else:
        pronoun_replacements = {}

    if brownText is not None:
        pronoun_replacements["[Brown]"] = brownText

    if os.path.exists(os.path.join(UPLOAD_FOLDER, 'wais.docx')):
        demoReps = extract_examinee_info(os.path.join(UPLOAD_FOLDER, 'wais.docx'))
        pronoun_replacements.update(demoReps)

    # Iterate through paragraphs and replace placeholders
    for para in doc.paragraphs:
        if "[First Name]" in para.text:
            para.text = para.text.replace("[First Name]", firstName)
        for placeholder, replacement in pronoun_replacements.items():
            if placeholder in para.text:
                para.text = para.text.replace(placeholder, replacement)

    # Iterate through tables (if any) and replace placeholders
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if "[First Name]" in para.text:
                        para.text = para.text.replace("[First Name]", firstName)
                    for placeholder, replacement in pronoun_replacements.items():
                        if placeholder in para.text:
                            para.text = para.text.replace(placeholder, replacement)

    # Save the modified document
    if os.name == 'posix':  # Unix/Linux/MacOS/Ubuntu
        save_path = doc.save(os.path.join(UPLOAD_FOLDER, doc_path))
    else:  # Windows, os.name will be 'nt'
        save_path = doc.save(doc_path)

