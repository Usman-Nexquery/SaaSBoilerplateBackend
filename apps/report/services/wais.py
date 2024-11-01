from pathlib import Path

from docx import Document
import re
import unicodedata
import os
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from .hi import get_indices
from config.settings import UPLOAD_FOLDER


def set_cell_vertical_alignment(cell, align="center"):
    """Set the vertical alignment of text within a cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    valign = OxmlElement('w:vAlign')
    valign.set(qn('w:val'), align)  # 'top', 'center', or 'bottom'
    tcPr.append(valign)

def extract_table_from_docx(docx_path):
    # Load the document
    doc = Document(docx_path)
    table_content = []  # List to hold rows of the table
    subtest_content = []

    rowLabels = ["Verbal Comprehension", "Perceptual Reasoning", "Working Memory", "Processing Speed",
                 "Full Scale"]

    subtestLabels = {
        "SI": "Similarities",
        "VC": "Vocabulary",
        "IN": "Information",
        "BD": "Block Design",
        "VP": "Visual Puzzles",
        "MR": "Matrix Reasoning",
        "FW": "Figure Weights",
        "DS": "Digit Span",
        "AR": "Arithmetic",
        "CD": "Coding",
        "SS": "Symbol Search"
    }
    reversed_subtestLabels = {value: key for key, value in subtestLabels.items()}
    table = {}

    #tableNums = [9, 11, 13, 15]
    tableNums = [1, 3, 4, 5, 6, 7]

    for i in range(len(tableNums)):
        # Extract data from the first table (for subtest_content)
        if len(doc.tables) > 0:
            table1 = doc.tables[tableNums[i]]
            for row in table1.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                subtest_content.append(row_data)


    # Extract data from the second table (for table_content)
    if len(doc.tables) > 1:
        table2 = doc.tables[4]
        for row in table2.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_content.append(row_data)
    for data in table_content:
        for label in rowLabels:
            print("======================================================== Data ====================================================")
            print(data[0])
            print(
                "======================================================== End ====================================================")
            if re.search(label, data[0], re.IGNORECASE) or 'Full Scale' in label:
                if label == "Verbal":
                    table["Verbal Comprehension"] = data
                elif 'Full Scale' in label:
                    table["Full Scale IQ"] = data
                    table["Full Scale IQ"][0] = 'Full Scale IQ'
                else:
                    table[label] = data
    print(table)
    # print('--------------------------------- SUBTEST_CONTENT -------------------------------')
    # print(subtest_content)
    for data in subtest_content:
        # print(f"This is for adding subtest_content: {data}")
        for label in subtestLabels:
            # print(f"data[1]: {data[1]}")
            if re.search(subtestLabels[label], data[0]) and subtestLabels[label] not in table:
                #if data[1].split('\n')[1] != subtestLabels[label] and data[1].split('\n')[0] != subtestLabels[label]:
                #    continue
#                data[0] = label
                data.append(label)
                table[subtestLabels[label]] = data  # Adjust the index based on how the data is structured


    return table

def normalize_string(s):
    s = re.sub(r'\([^)]*\)', '', s)
    # print('s: ', s)
    if 'Full Scale' in s:
        s = 'Full Scale IQ'
        return s
    return unicodedata.normalize("NFKC", s.strip())

def normalize_key(key):
    # Remove content within parentheses and extra spaces
    return ' '.join(key.split()[:-1]).strip()

def insert_table_into_word(path, data, filename):
    # Load the Word document
    doc = Document(path)

    table_indices = get_indices(doc, "WAIS")
    # print(table_indices)

    wais_index = [3,4,5]
    wais_subtest = [2, 3] # This is the list position of Scaled Score, Percentile

    # Process each table found by the getIndices function
    for table_name, info in table_indices.items():
        print("Table name: ", table_name)
        print("info: ", info)
        table = doc.tables[info["table_index"]]  # Access the specific table by index
        column_indices = info["columns"]  # Get the list of column indices needed
        # column_indices = [x + 1 for x in column_indices]
        # Iterate over each row in the table after the header row
        #print(data)
        for row in table.rows[1:]:  # Skipping the header row
            key = row.cells[0].text.strip()
            #print("Key: ", key)

            if key in data or normalize_key(key) in data:
                # Insert data into specified columns
                for i, column_index in enumerate(column_indices):
                    cell = row.cells[column_index]
                    #if len(key.split(" ")) > 2:
                    if table_name == "WAIS-V Index":
                        if 'Verbal Comprehension Index' in normalize_key(key):
                            key = normalize_key(key)
                        print(key)
                        print('Data: ', data[normalize_key(key)])
                        cell.text = str(data[normalize_key(key)][wais_index[i]])  # Assuming data[key] is a list with the needed data

                    else:
                        while 'Index' in key:
                            key = normalize_key(key)
                        if 'Full Scale IQ' in data[key][0]:
                            continue
                        print('---------------------------------')
                        print(key)
                        print('Data: ', data[key])
                        cell.text = str(data[key][wais_subtest[i]])  # Assuming data[key] is a list with the needed data
                        print(cell.text)
                        print('---------------------------------')
                    paragraph = cell.paragraphs[0]
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    # Assume set_cell_vertical_alignment is defined elsewhere
                    set_cell_vertical_alignment(cell, "center")

    # Save the document

    if os.name == 'posix':  # Unix/Linux/MacOS/Ubuntu
        doc.save(os.path.join(UPLOAD_FOLDER, filename))
    else:  # Windows, os.name will be 'nt'
        doc.save(filename)
def ordinal(number):
    """Convert an integer into its ordinal representation."""
    suffix = ['th', 'st', 'nd', 'rd', 'th'][min(number % 10, 4)]
    if 11 <= (number % 100) <= 13:
        suffix = 'th'
    return f"{number}{suffix}"
def find_key_by_value(data, search_value):
    """Search through dictionary to find a key based on a value substring."""
    for key, values in data.items():
        #if any(search_value in str(val) for val in values):
        #    return key
        if any(re.search(r"\b" + re.escape(search_value) + r"\b", val) for val in values):
            return key
    return None
def update_document(template_path, data, filename):
    doc = Document(template_path)

    # Define patterns to look for specific placeholder texts
    placeholders = {
        'VCI_DESC': ('[VCI_DESC]', 6),  # 7th item in the list for description
        'VCI_PCT': ('[VCI_PCT]', 4),  # 5th item in the list for percentile
        'VCI_BOLD': ('[VCI_BOLD]', 2),
        'VCI_LONG_BOLD': ('[VCI_LONG_BOLD]', 0),
        'SI_UL': ('[SI_UL]', 0),
        'VC_UL': ('[VC_UL]', 0),
        'IN_UL': ('[IN_UL]', 0),
        'PRI_BOLD': ('[PRI_BOLD]', 2),
        'PRI_LONG_BOLD': ('[PRI_LONG_BOLD]', 0),
        'PRI_DESC': ('[PRI_DESC]', 6),  # 7th item in the list for description
        'PRI_PCT': ('[PRI_PCT]', 4),  # 5th item in the list for percentile
        'MR_UL': ('[MR_UL]', 0),
        'BD_UL': ('[BD_UL]', 0),
        'VP_UL': ('[VP_UL]', 0),
        'WMI_BOLD': ('[WMI_BOLD]', 2),
        'WMI_LONG_BOLD': ('[WMI_LONG_BOLD]', 0),
        'WMI_DESC': ('[WMI_DESC]', 6),  # 7th item in the list for description
        'WMI_PCT': ('[WMI_PCT]', 4),  # 5th item in the list for percentile
        'DS_UL': ('[DS_UL]', 0),
        'AR_UL': ('[AR_UL]', 0),
        'PSI_BOLD': ('[PSI_BOLD]', 2),
        'PSI_LONG_BOLD': ('[PSI_LONG_BOLD]', 0),
        'PSI_DESC': ('[PSI_DESC]', 6),  # 7th item in the list for description
        'PSI_PCT': ('[PSI_PCT]', 4),  # 5th item in the list for percentile
        'SS_UL': ('[SS_UL]', 0),
        'CD_UL': ('[CD_UL]', 0)

    }
    # print('================================================================ here is the data =======================================================================')
    # print(data)

    for paragraph in doc.paragraphs:
        original_text = paragraph.text
        changes = []
        # Find all placeholders and prepare changes
        for key, (placeholder, idx) in placeholders.items():
            abbreviation = key.split('_')[0]
            full_key = find_key_by_value(data, abbreviation)
            if abbreviation == "SI":
                pass
                # print('_________________________________________________________________________________')
                # print(data)
                # print(f"full_key: {full_key}")
                #
                # print(f"placeholders: {placeholder}")
                # print(f"original_text: {original_text}")

            if full_key and placeholder in original_text:
                # print(f"key: {key}")
                # print(data)
                # print(f"full_key: {full_key}")
                # print(f"idx: {idx}")
                description = data[full_key][idx]
                # print(f"key: {key}")
                if 'PCT' in key:
                    percentile_value = int(description)
                    description = f"{ordinal(percentile_value)} percentile"
                    tformat = 'italic'
                elif 'VCI_LONG_BOLD' in key:
                    description = 'Verbal Comprehension Index (VCI)'
                    tformat = 'bold'
                elif 'VCI_BOLD' in key:
                    description = 'VCI'
                    tformat = 'bold'
                elif 'PRI_LONG_BOLD' in key:
                    description = 'Perceptual Reasoning Index (PRI)'
                    tformat = 'bold'
                # elif 'VSI_BOLD' in key:
                #     description = 'VSI'
                #     tformat = 'bold'
                elif 'WMI_LONG_BOLD' in key:
                    description = 'Working Memory Index (WMI)'
                    tformat = 'bold'
                elif 'PSI_LONG_BOLD' in key:
                    description = 'Processing Speed Index (PSI)'
                    tformat = 'bold'
                elif 'BOLD' in key:
                    tformat = 'bold'
                    # print(f"full_key: {full_key}")
                    # print(f"placeholder: {placeholder}")
                    # print(f"abbreviation: {abbreviation}")
                elif 'UL' in key:
                    tformat = 'underline'
                    # print(f"full_key: {full_key}")
                    # print(f"placeholder: {placeholder}")
                    # print(f"abbreviation: {abbreviation}")
                else:
                    description = data[full_key][idx].lower()
                    tformat = 'italic'
                for m in re.finditer(re.escape(placeholder), original_text):
                    changes.append((m.start(), m.end(), description, tformat))
        #print('changes:', changes)
        if changes:
            # Sort changes based on the start index
            changes.sort()
            # Split the paragraph and apply changes
            last_idx = 0
            paragraph.clear()
            for start, end, text, tformat in changes:
                if tformat == 'italic':
                    paragraph.add_run(original_text[last_idx:start])  # Add text before placeholder
                    italic_run = paragraph.add_run(text)
                    italic_run.italic = True  # Apply italic formatting to replacement
                    last_idx = end
                if tformat == 'bold':
                    paragraph.add_run(original_text[last_idx:start])  # Add text before placeholder
                    bold_run = paragraph.add_run(text)
                    bold_run.bold = True  # Apply italic formatting to replacement
                    last_idx = end
                elif tformat == 'underline':
                    paragraph.add_run(original_text[last_idx:start])  # Add text before placeholder
                    ul_run = paragraph.add_run(text)
                    ul_run.underline = True  # Apply italic formatting to replacement
                    last_idx = end
            paragraph.add_run(original_text[last_idx:])  # Add remaining text after last placeholder

    # Save the modified document
    if os.name == 'posix':  # Unix/Linux/MacOS/Ubuntu
        save_path = doc.save(os.path.join(UPLOAD_FOLDER, filename))
    else:  # Windows, os.name will be 'nt'
        save_path = doc.save(filename)

# pdf_path = 'wais.doc'
# doc_path = 'template.docx'
# filename = 'report.docx'
# #
# # Extract table data from PDF
# table_content = extract_table_from_docx(pdf_path)
# #
# #
# # Insert table data into Word document
# insert_table_into_word(doc_path, table_content, filename)
