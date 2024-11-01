
from docx import Document
import re
import unicodedata
import os
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from config.settings import UPLOAD_FOLDER

def set_cell_vertical_alignment(cell, align="center"):
    """Set the vertical alignment of text within a cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    valign = OxmlElement('w:vAlign')
    valign.set(qn('w:val'), align)  # 'top', 'center', or 'bottom'
    tcPr.append(valign)

def extract_table_from_docx(docx_path):
    # Load the document
    doc = Document(docx_path)
    table_content = []  # List to hold rows of the table
    subtest_content = []

    rowLabels = ["Verbal", "Visual Spatial", "Fluid Reasoning", "Working Memory", "Processing Speed",
                 "Full Scale IQ"]

    subtestLabels = {
        "SI": "Similarities",
        "VC": "Vocabulary",
        "BD": "Block Design",
        "VP": "Visual Puzzles",
        "MR": "Matrix Reasoning",
        "FW": "Figure Weights",
        "DS": "Digit Span",
        "PS": "Picture Span",
        "CD": "Coding",
        "SS": "Symbol Search"
    }

    table = {}

    # Extract data from the first table (for subtest_content)
    if len(doc.tables) > 0:
        table1 = doc.tables[3]
        for row in table1.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            subtest_content.append(row_data)
    # Extract data from the second table (for table_content)
    if len(doc.tables) > 1:
        table2 = doc.tables[5]
        for row in table2.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_content.append(row_data)
    for data in table_content:
        for label in rowLabels:
            if re.search(label, data[0], re.IGNORECASE):
                if label == "Verbal":
                    table["Verbal Comprehension"] = data
                else:
                    table[label] = data

    for data in subtest_content:
        for label in subtestLabels:
            if re.search(subtestLabels[label], data[1]) and subtestLabels[label] not in table:
                #if data[1].split('\n')[1] != subtestLabels[label] and data[1].split('\n')[0] != subtestLabels[label]:
                #    continue
                table[subtestLabels[label]] = data[1:]  # Adjust the index based on how the data is structured

    return table
def extract_table_from_pdf(pdf_path):
    doc = fitz.open(pdf_path)
    table_content = []  # List to hold rows of the table
    subtest_content = []

    rowLabels = ["Verbal Comprehension", "Visual Spatial", "Fluid Reasoning", "Working Memory", "Processing Speed",
                 "Full Scale IQ"]

    subtestLabels = {
        "SI": "Similarities",
        "VC": "Vocabulary",
        "BD": "Block Design",
        "VP": "Visual Puzzles",
        "MR": "Matrix Reasoning",
        "FW": "Figure Weights",
        "DS": "Digit Span",
        "PS": "Picture Span",
        "CD": "Coding",
        "SS": "Symbol Search"
    }

    table = {}
    for page in doc:
        if page.number == 2:
            blocks = page.get_text("blocks")
            blocks.sort(key=lambda block: (block[1], block[0]))  # Sort text blocks
            current_row = []
            current_y = None
            for block in blocks:
                # Block format: (x0, y0, x1, y1, "text", block_type, block_number)
                x0, y0, text = block[0], block[1], block[4]
                if current_y is None or abs(y0 - current_y) > 5:  # New row
                    if current_row:  # Save previous row
                        table_content.append(current_row)
                    current_row = [text]
                    current_y = y0
                else:
                    current_row.append(text)
            if current_row:  # Save the last row
                table_content.append(current_row)
        if page.number == 1:
            blocks = page.get_text("blocks")
            blocks.sort(key=lambda block: (block[1], block[0]))  # Sort text blocks
            current_row = []
            current_y = None
            for block in blocks:
                # Block format: (x0, y0, x1, y1, "text", block_type, block_number)
                x0, y0, text = block[0], block[1], block[4]
                if current_y is None or abs(y0 - current_y) > 5:  # New row
                    if current_row:  # Save previous row
                        subtest_content.append(current_row)
                    current_row = [text]
                    current_y = y0
                else:
                    current_row.append(text)
            if current_row:  # Save the last row
                subtest_content.append(current_row)

    for data in table_content:
        for label in rowLabels:
            if re.search(label, data[0], re.IGNORECASE):
                table[label] = data[0].split('\n')

    for data in subtest_content:
        for label in subtestLabels:
            if re.search(subtestLabels[label], data[0]) and subtestLabels[label] not in table:
                if data[0].split('\n')[1] != subtestLabels[label] and data[0].split('\n')[0] != subtestLabels[label]:
                    continue
                table[subtestLabels[label]] = data[0].split('\n')[1:]
    #print(table)
    return table


def normalize_string(s):
    s = re.sub(r'\([^)]*\)', '', s)
    return unicodedata.normalize("NFKC", s.strip())


def insert_table_into_word(path, data, filename):
    # Load the Word document
    doc = Document(path)

    # Assume the first table is the target
    table = doc.tables[0]

    # Iterate over each row in the table after the header row
    for row in table.rows[1:]:  # Skipping the header row

        key = row.cells[0].text
        key = normalize_string(key)
        print(key)
        if key in data:
            # Insert the corresponding 4th, 5th, and 7th items from the list into the 2nd, 3rd, and 4th columns

            for i, content_index in zip(range(1, 4), [3, 4, 6]):
                cell = row.cells[i]
                cell.text = data[key][content_index]
                paragraph = cell.paragraphs[0]
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_cell_vertical_alignment(cell, "center")

    # Assume the first table is the target
    table = doc.tables[1]

    # Iterate over each row in the table after the header row
    for row in table.rows[1:]:  # Skipping the header row
        key = row.cells[0].text
        key = normalize_string(key)

        if key in data:
            # Insert the corresponding 4th, 5th, and 7th items from the list into the 2nd, 3rd, and 4th columns
            for i, content_index in zip(range(1, 3), [3, 4]):
                if 'Puzzles' in key or 'Weight' in key or 'Picture' in key or 'Symbol' in key:
                    cell = row.cells[i]
                    cell.text = data[key][content_index]
                    paragraph = cell.paragraphs[0]
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    set_cell_vertical_alignment(cell, "center")
                else:
                    cell = row.cells[i]
                    cell.text = data[key][content_index]
                    paragraph = cell.paragraphs[0]
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    set_cell_vertical_alignment(cell, "center")

    # Save the document

    if os.name == 'posix':  # Unix/Linux/MacOS/Ubuntu
        save_path = doc.save(os.path.join(UPLOAD_FOLDER, filename))
    else:  # Windows, os.name will be 'nt'
        save_path = doc.save('updated_template.docx')

def ordinal(number):
    """Convert an integer into its ordinal representation."""
    suffix = ['th', 'st', 'nd', 'rd', 'th'][min(number % 10, 4)]
    if 11 <= (number % 100) <= 13:
        suffix = 'th'
    return f"{number}{suffix}"
def find_key_by_value(data, search_value):
    """Search through dictionary to find a key based on a value substring."""
    for key, values in data.items():
        #if any(search_value in str(val) for val in values):
        #    return key
        if any(re.search(r"\b" + re.escape(search_value) + r"\b", val) for val in values):
            return key
    return None
def update_document(template_path, data, filename):
    doc = Document(template_path)

    # Define patterns to look for specific placeholder texts
    placeholders = {
        'VCI_DESC': ('[VCI_DESC]', 6),  # 7th item in the list for description
        'VCI_PCT': ('[VCI_PCT]', 4),  # 5th item in the list for percentile
        'VSI_DESC': ('[VSI_DESC]', 6),
        'VSI_PCT': ('[VSI_PCT]', 4),
        'FRI_DESC': ('[FRI_DESC]', 6),
        'FRI_PCT': ('[FRI_PCT]', 4),
        'WMI_DESC': ('[WMI_DESC]', 6),
        'WMI_PCT': ('[WMI_PCT]', 4),
        'PSI_DESC': ('[PSI_DESC]', 6),
        'PSI_PCT': ('[PSI_PCT]', 4),
        'VCI_BOLD': ('[VCI_BOLD]', 1),
        'VCI_LONG_BOLD': ('[VCI_LONG_BOLD]', 0),
        'SI_UL': ('[SI_UL]', 0),
        'VC_UL': ('[VC_UL]', 0),
        'VSI_BOLD': ('[VSI_BOLD]', 1),
        'VSI_LONG_BOLD': ('[VSI_LONG_BOLD]', 0),
        'BD_UL': ('[BD_UL]', 0),
        'VP_UL': ('[VP_UL]', 0),
        'FRI_BOLD': ('[FRI_BOLD]', 1),
        'FRI_LONG_BOLD': ('[FRI_LONG_BOLD]', 0),
        'MR_UL': ('[MR_UL]', 0),
        'FW_UL': ('[FW_UL]', 0),
        'WMI_BOLD': ('[WMI_BOLD]', 1),
        'WMI_LONG_BOLD': ('[WMI_LONG_BOLD]', 0),
        'PS_UL': ('[PS_UL]', 0),
        'DS_UL': ('[DS_UL]', 0),
        'PSI_BOLD': ('[PSI_BOLD]', 1),
        'PSI_LONG_BOLD': ('[PSI_LONG_BOLD]', 0),
        'SS_UL': ('[SS_UL]', 0),
        'CD_UL': ('[CD_UL]', 0)

    }
    print('here is the data')
    print(data)

    for paragraph in doc.paragraphs:
        original_text = paragraph.text
        changes = []
        # Find all placeholders and prepare changes
        for key, (placeholder, idx) in placeholders.items():
            abbreviation = key.split('_')[0]
            full_key = find_key_by_value(data, abbreviation)

            if full_key and placeholder in original_text:
                description = data[full_key][idx]
                print(f"key: {key}")
                if 'PCT' in key:
                    percentile_value = int(description)
                    description = f"{ordinal(percentile_value)} percentile"
                    tformat = 'italic'
                elif 'VCI_LONG_BOLD' in key:
                    description = 'Verbal Comprehension Index (VCI)'
                    tformat = 'bold'
                elif 'VCI_BOLD' in key:
                    description = 'VCI'
                    tformat = 'bold'
                elif 'VSI_LONG_BOLD' in key:
                    description = 'Visual Spatial Index (VSI)'
                    tformat = 'bold'
                # elif 'VSI_BOLD' in key:
                #     description = 'VSI'
                #     tformat = 'bold'
                elif 'FRI_LONG_BOLD' in key:
                    description = 'Fluid Reasoning Index (FRI)'
                    tformat = 'bold'
                elif 'WMI_LONG_BOLD' in key:
                    description = 'Working Memory Index (WMI)'
                    tformat = 'bold'
                elif 'PSI_LONG_BOLD' in key:
                    description = 'Processing Speed Index (PSI)'
                    tformat = 'bold'
                elif 'BOLD' in key:
                    tformat = 'bold'
                    print(f"full_key: {full_key}")
                    print(f"placeholder: {placeholder}")
                    print(f"abbreviation: {abbreviation}")
                elif 'UL' in key:
                    tformat = 'underline'
                    print(f"full_key: {full_key}")
                    print(f"placeholder: {placeholder}")
                    print(f"abbreviation: {abbreviation}")
                else:
                    description = data[full_key][idx].lower()
                    tformat = 'italic'
                for m in re.finditer(re.escape(placeholder), original_text):
                    changes.append((m.start(), m.end(), description, tformat))
        print('changes:', changes)
        if changes:
            # Sort changes based on the start index
            changes.sort()
            # Split the paragraph and apply changes
            last_idx = 0
            paragraph.clear()
            for start, end, text, tformat in changes:
                if tformat == 'italic':
                    paragraph.add_run(original_text[last_idx:start])  # Add text before placeholder
                    italic_run = paragraph.add_run(text)
                    italic_run.italic = True  # Apply italic formatting to replacement
                    last_idx = end
                if tformat == 'bold':
                    paragraph.add_run(original_text[last_idx:start])  # Add text before placeholder
                    bold_run = paragraph.add_run(text)
                    bold_run.bold = True  # Apply italic formatting to replacement
                    last_idx = end
                elif tformat == 'underline':
                    paragraph.add_run(original_text[last_idx:start])  # Add text before placeholder
                    ul_run = paragraph.add_run(text)
                    ul_run.underline = True  # Apply italic formatting to replacement
                    last_idx = end
            paragraph.add_run(original_text[last_idx:])  # Add remaining text after last placeholder

    # Save the modified document
    if os.name == 'posix':  # Unix/Linux/MacOS/Ubuntu
        save_path = doc.save(os.path.join(UPLOAD_FOLDER, filename))
    else:  # Windows, os.name will be 'nt'
        save_path = doc.save(filename)

    #update_document(filename, data, filename)
# pdf_path = 'wisc.pdf'
# docx_path = 'wisc.docx'
# doc_path = 'template.docx'
# filename = 'report.docx'
#
# table_content = extract_table_from_pdf(pdf_path)
# print("pdf_content", table_content)
#
# # Extract table data from PDF
# table_content = extract_table_from_docx(docx_path)
# print("docx_content", table_content)
# # Insert table data into Word document
# insert_table_into_word(doc_path, table_content, filename)
# print('inserting')
# update_document(doc_path, table_content, filename)
# print('done')
