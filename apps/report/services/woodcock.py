from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os
import re
from config.settings import UPLOAD_FOLDER


def set_cell_vertical_alignment(cell, align="center"):
    """Set the vertical alignment of text within a cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    valign = OxmlElement('w:vAlign')
    valign.set(qn('w:val'), align)  # 'top', 'center', or 'bottom'
    tcPr.append(valign)



def extract_table_data(filepath):
    doc = Document(filepath)
    table = doc.tables[1]  # Assuming the third table contains the data
    data = {}
    row = []
    collect = False
    i = 1


    if not table._element.xpath('.//w:tblGrid'):
        tblGrid = OxmlElement('w:tblGrid')
        table._element.insert(-1, tblGrid)


    for cell in table._cells[5:]:

        if cell.text.startswith("READING"):
            collect = True  # Start collecting text
        if collect:
            row.append(cell.text)
            print(row)
        if i%5 == 0:
            data[row[0]] = row
            row = []
        i += 1


    return data

def insert_table_into_word(path, data, filename):
    # Load the Word document
    doc = Document(path)

    # Assume the first table is the target
    table = doc.tables[4]

    # Iterate over each row in the table after the header row
    for row in table.rows[1:]:  # Skipping the header row
        key = row.cells[0].text
        print("key", key)
        print("data", data)
        if key in data:
            # Insert the corresponding 4th, 5th, and 7th items from the list into the 2nd, 3rd, and 4th columns
            for i, content_index in zip(range(1, 4), [1, 3, 4]):
                cell = row.cells[i]
                cell.text = data[key][content_index]
                paragraph = cell.paragraphs[0]
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_cell_vertical_alignment(cell, "center")



    # Save the document
    #doc.save(os.path.join(UPLOAD_FOLDER, filename))
    #doc.save(filename)
    if os.name == 'posix':  # Unix/Linux/MacOS/Ubuntu
        save_path = doc.save(os.path.join(UPLOAD_FOLDER, filename))
    else:  # Windows, os.name will be 'nt'
        save_path = doc.save(filename)
def ordinal(number):
    """Convert an integer into its ordinal representation."""
    suffix = ['th', 'st', 'nd', 'rd', 'th'][min(number % 10, 4)]
    if 11 <= (number % 100) <= 13:
        suffix = 'th'
    return f"{number}{suffix}"
def find_key_by_value(data, search_value):
    """Search through dictionary to find a key based on a value substring."""
    for key, values in data.items():
        #if any(search_value in str(val) for val in values):
        #    return key
        if any(re.search(r"\b" + re.escape(search_value) + r"\b", val) for val in values):
            return key
    return None
def score_to_category(score):
    if score > 130:
        return "very superior"
    elif 120 <= score <= 129:
        return "superior"
    elif 110 <= score <= 119:
        return "high average"
    elif 90 <= score <= 109:
        return "average"
    elif 80 <= score <= 89:
        return "low average"
    elif 70 <= score <= 79:
        return "low"
    else:
        return "very low"

def set_font(run):
    # Set the font to Times New Roman for the run
    run.font.name = 'Times New Roman'

    # Ensure compatibility with Word
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    rPr.append(rFonts)
    return run
def update_document(template_path, data, filename):
    doc = Document(template_path)

    # Define patterns to look for specific placeholder texts
    placeholders = {
        'BR_DESC': ('[BR_DESC]', 3),  # 7th item in the list for description
        'BR_PCT': ('[BR_PCT]', 4),  # 5th item in the list for percentile
        'BR_BOLD': ('[BR_BOLD]', 0),
        'BM_DESC': ('[BM_DESC]', 3),
        'BM_PCT': ('[BM_PCT]', 4),
        'BM_BOLD': ('[BM_BOLD]', 0),
        'BRS_DESC': ('[BRS_DESC]', 3),
        'BRS_PCT': ('[BRS_PCT]', 4),
        'BRS_BOLD': ('[BRS_BOLD]', 0),
        'BW_DESC': ('[BW_DESC]', 3),
        'BW_PCT': ('[BW_PCT]', 4),
        'BW_BOLD': ('[BW_BOLD]', 0),
        'RC_DESC': ('[RC_DESC]', 3),
        'RC_PCT': ('[RC_PCT]', 4),
        'RC_BOLD': ('[RC_BOLD]', 0),
        'LW_UL': ('[LW_UL]', 0),
        'LW_PCT': ('[LW_PCT]', 4),
        'PC_UL': ('[PC_UL]', 0),
        'PC_PCT': ('[PC_PCT]', 4),
        'SR_UL': ('[SR_UL]', 0),
        'SR_PCT': ('[SR_PCT]', 4),
        'WA_UL': ('[WA_UL]', 0),
        'WA_PCT': ('[WA_PCT]', 4),
        'AP_UL': ('[AP_UL]', 0),
        'AP_PCT': ('[AP_PCT]', 4),
        'CA_UL': ('[CA_UL]', 0),
        'CA_PCT': ('[CA_PCT]', 4),
        'MF_UL': ('[MF_UL]', 0),
        'MF_PCT': ('[MF_PCT]', 4),
        'SP_UL': ('[SP_UL]', 0),
        'SP_PCT': ('[SP_PCT]', 4),
        'WS_UL': ('[WS_UL]', 0),
        'WS_PCT': ('[WS_PCT]', 4),
        'WF_UL': ('[WF_UL]', 0),
        'WF_PCT': ('[WF_PCT]', 4),
        'SW_UL': ('[SW_UL]', 0),
        'SW_PCT': ('[SW_PCT]', 4),
        'RR_UL': ('[RR_UL]', 0),
        'RR_PCT': ('[RR_PCT]', 4)
    }

    glossary = {
        'BR': 'BROAD READING',
        'SR': 'Sentence Reading Fluency',
        'LW': 'Letter-Word Identification',
        'PC': 'Passage Comprehension',
        'BRS': 'BASIC READING SKILLS',
        'WA': 'Word Attack',
        'AP': 'Applied Problems',
        'CA': 'Calculation',
        'BM': 'BROAD MATHEMATICS',
        'MF': 'Math Facts Fluency',
        'SP': 'Spelling',
        'WS': 'Writing Samples',
        'BW': 'BROAD WRITTEN LANGUAGE',
        'SW': 'Sentence Writing Fluency',
        'WF': 'Sentence Writing Fluency'
    }
    print('here is the data')
    print(data)

    for paragraph in doc.paragraphs:
        original_text = paragraph.text
        changes = []
        # Find all placeholders and prepare changes
        for key, (placeholder, idx) in placeholders.items():
            abbreviation = glossary[key.split('_')[0]]
            full_key = find_key_by_value(data, abbreviation)
            print(f"abbreviation {abbreviation}")
            print(f"full_key {full_key}")
            if full_key and placeholder in original_text:
                description = data[full_key][idx]
                print(f"key: {key}")
                if 'PCT' in key:
                    if '<' in description or float(description) < 1:
                        percentile_value = description
                        description = f"{percentile_value} percentile"
                    else:
                        percentile_value = int(description)
                        description = f"{ordinal(percentile_value)} percentile"
                    tformat = 'italic'
                elif 'BOLD' in key:
                    tformat = 'bold'
                elif 'UL' in key:
                    tformat = 'underline'
                else:
                    description = score_to_category(int(data[full_key][idx]))
                    tformat = 'italic'
                for m in re.finditer(re.escape(placeholder), original_text):
                    changes.append((m.start(), m.end(), description, tformat))
        print('changes:', changes)
        if changes:
            # Sort changes based on the start index
            changes.sort()
            # Split the paragraph and apply changes
            last_idx = 0
            paragraph.clear()
            for start, end, text, tformat in changes:
                if tformat == 'italic':
                    run = paragraph.add_run(original_text[last_idx:start])  # Add text before placeholder

                    set_font(run)

                    italic_run = paragraph.add_run(text.strip())
                    set_font(italic_run)
                    italic_run.italic = True  # Apply italic formatting to replacement

                    last_idx = end
                if tformat == 'bold':
                    run = paragraph.add_run(original_text[last_idx:start])  # Add text before placeholder
                    set_font(run)

                    bold_run = paragraph.add_run(text.title().strip())
                    set_font(bold_run)
                    bold_run.bold = True  # Apply italic formatting to replacement
                    last_idx = end
                elif tformat == 'underline':
                    run = paragraph.add_run(original_text[last_idx:start])  # Add text before placeholder
                    set_font(run)

                    ul_run = paragraph.add_run(text.title().strip())
                    set_font(ul_run)
                    ul_run.underline = True  # Apply italic formatting to replacement
                    last_idx = end
            run = paragraph.add_run(original_text[last_idx:])  # Add remaining text after last placeholder
            set_font(run)

    # Save the modified document
    if os.name == 'posix':  # Unix/Linux/MacOS/Ubuntu
        save_path = doc.save(os.path.join(UPLOAD_FOLDER, filename))
    else:  # Windows, os.name will be 'nt'
        save_path = doc.save(filename)

# Adjust the file paths as necessary
#source_filepath = 'woodcock.docx'
#template_filepath = 'template.docx'
#filename = 'report.docx'
#data = extract_table_data(source_filepath)
#insert_table_into_word(template_filepath, data, filename)
